"""
Resume exporters for PDF, DOCX, and TXT formats
"""

import logging
from pathlib import Path
from typing import Optional

from .base import BaseExporter

logger = logging.getLogger(__name__)


class TextExporter(BaseExporter):
    """Export resume as plain text"""

    def __init__(self):
        super().__init__("txt")

    async def export(
        self, resume_text: str, output_path: str, metadata: dict = None
    ) -> bool:
        """Export to plain text file"""
        try:
            Path(output_path).write_text(resume_text, encoding="utf-8")
            self.logger.info(f"Exported to {output_path}")
            return True
        except Exception as e:
            self.logger.error(f"Export failed: {e}")
            return False


class PDFExporter(BaseExporter):
    """Export resume as PDF"""

    def __init__(self):
        super().__init__("pdf")

    async def export(
        self, resume_text: str, output_path: str, metadata: dict = None
    ) -> bool:
        """Export to PDF file"""
        try:
            # Use reportlab for PDF generation
            from reportlab.lib.pagesizes import letter
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
            from reportlab.lib.units import inch

            doc = SimpleDocTemplate(
                output_path,
                pagesize=letter,
                rightMargin=0.5 * inch,
                leftMargin=0.5 * inch,
                topMargin=0.5 * inch,
                bottomMargin=0.5 * inch,
            )

            # Build PDF content
            story = []
            styles = getSampleStyleSheet()

            # Custom styles
            title_style = ParagraphStyle(
                "CustomTitle",
                parent=styles["Heading1"],
                fontSize=14,
                textColor="#000000",
                spaceAfter=6,
                alignment=1,  # Center
            )

            body_style = ParagraphStyle(
                "CustomBody",
                parent=styles["BodyText"],
                fontSize=11,
                leading=14,
            )

            # Parse and add resume content
            for line in resume_text.split("\n"):
                if line.strip():
                    # Detect headers (all caps, short lines)
                    if line.isupper() and len(line) < 50:
                        story.append(Paragraph(line, title_style))
                    else:
                        story.append(Paragraph(line, body_style))
                else:
                    story.append(Spacer(1, 0.1 * inch))

            # Build PDF
            doc.build(story)
            self.logger.info(f"Exported to {output_path}")
            return True

        except ImportError:
            self.logger.error("reportlab not installed: pip install reportlab")
            return False
        except Exception as e:
            self.logger.error(f"PDF export failed: {e}")
            return False


class DOCXExporter(BaseExporter):
    """Export resume as DOCX (Word)"""

    def __init__(self):
        super().__init__("docx")

    async def export(
        self, resume_text: str, output_path: str, metadata: dict = None
    ) -> bool:
        """Export to DOCX file"""
        try:
            from docx import Document
            from docx.shared import Pt, Inches, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            doc = Document()

            # Set margins
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(0.5)
                section.bottom_margin = Inches(0.5)
                section.left_margin = Inches(0.5)
                section.right_margin = Inches(0.5)

            # Parse resume content
            lines = resume_text.split("\n")

            for line in lines:
                if not line.strip():
                    doc.add_paragraph("")
                    continue

                # Check if it's a section header
                if line.isupper() and len(line) < 50:
                    # Section header
                    p = doc.add_paragraph(line)
                    p.runs[0].font.bold = True
                    p.runs[0].font.size = Pt(12)
                    p.paragraph_format.space_before = Pt(12)
                    p.paragraph_format.space_after = Pt(6)
                elif line.startswith("  •"):
                    # Bullet point
                    p = doc.add_paragraph(line.strip(), style="List Bullet")
                    p.paragraph_format.left_indent = Inches(0.25)
                else:
                    # Regular paragraph
                    doc.add_paragraph(line)

            doc.save(output_path)
            self.logger.info(f"Exported to {output_path}")
            return True

        except ImportError:
            self.logger.error("python-docx not installed: pip install python-docx")
            return False
        except Exception as e:
            self.logger.error(f"DOCX export failed: {e}")
            return False


class ExportManager:
    """Manage resume exports across formats"""

    def __init__(self):
        self.exporters = {
            "txt": TextExporter(),
            "pdf": PDFExporter(),
            "docx": DOCXExporter(),
        }

    async def export(
        self, resume_text: str, output_path: str, format_type: str, metadata: dict = None
    ) -> tuple[bool, str]:
        """
        Export resume in specified format.

        Args:
            resume_text: Formatted resume text
            output_path: Path to output file
            format_type: 'txt', 'pdf', or 'docx'
            metadata: Optional metadata

        Returns:
            Tuple of (success, message)
        """
        if format_type not in self.exporters:
            return False, f"Unsupported format: {format_type}. Use txt, pdf, or docx."

        exporter = self.exporters[format_type]
        success = await exporter.export(resume_text, output_path, metadata)

        if success:
            return True, f"Resume exported successfully to {output_path}"
        else:
            return False, f"Failed to export resume as {format_type}"

    async def export_all(
        self, resume_text: str, output_dir: str, metadata: dict = None
    ) -> dict:
        """
        Export resume in all formats.

        Returns:
            Dict with results for each format
        """
        results = {}

        for format_type in ["txt", "pdf", "docx"]:
            filename = metadata.get("filename", "resume") if metadata else "resume"
            output_path = f"{output_dir}/{filename}.{format_type}"

            success, message = await self.export(resume_text, output_path, format_type, metadata)
            results[format_type] = {"success": success, "message": message, "path": output_path}

        return results
